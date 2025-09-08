import os
import io
import re
from datetime import datetime
from mimetypes import guess_type
from flask import (
    Flask, render_template, request, redirect,
    url_for, session, flash, send_file
)
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import or_, func
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from docx import Document
from flask_mail import Mail, Message
from config import Config

# ----------------- НАСТРОЙКИ -----------------
app = Flask(__name__)
app.config.from_object(Config)

db = SQLAlchemy(app)
mail = Mail(app)

# ----------------- МОДЕЛИ -----------------
class User(db.Model):
    __tablename__ = 'user'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password_hash = db.Column(db.String(250), nullable=True)
    email = db.Column(db.String(120), unique=True, nullable=False)
    full_name = db.Column(db.String(120), nullable=True)
    avatar_url = db.Column(db.String(250), nullable=True)
    progress_percent = db.Column(db.Integer, default=0, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow, nullable=False)

    def set_password(self, raw_password: str):
        self.password_hash = generate_password_hash(raw_password)

    def check_password(self, raw_password: str) -> bool:
        if not self.password_hash:
            return False
        return check_password_hash(self.password_hash, raw_password)


class Material(db.Model):
    __tablename__ = 'material'
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(512), nullable=False)
    file_data = db.Column(db.LargeBinary, nullable=True)
    file_name = db.Column(db.String(255), nullable=True)
    type = db.Column(db.String(20), nullable=False, default='theory')  # theory/practice
    language = db.Column(db.String(50), nullable=False, default='python')
    created_at = db.Column(db.DateTime, server_default=func.now(), nullable=False)


class MaterialOpen(db.Model):
    __tablename__ = 'material_open'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id', ondelete='CASCADE'), nullable=False)
    material_id = db.Column(db.Integer, db.ForeignKey('material.id', ondelete='CASCADE'), nullable=False)
    opened_at = db.Column(db.DateTime, server_default=func.now(), nullable=False)

    user = db.relationship('User')
    material = db.relationship('Material')


with app.app_context():
    db.create_all()

# ----------------- ХЕЛПЕРЫ -----------------
USERNAME_RE = re.compile(r'^[A-Za-z0-9_]{3,30}$')
EMAIL_RE = re.compile(r'^[^@]+@[^@]+\.[^@]+$')
PWD_RE = re.compile(r'^(?=.*[a-z])(?=.*[A-Z])(?=.*\d).{8,}$')

def log_open(material_id: int):
    uid = session.get('user_id')
    if not uid:
        return
    try:
        db.session.add(MaterialOpen(user_id=uid, material_id=material_id))
        db.session.commit()
    except Exception:
        db.session.rollback()

# ----------------- ГЛАВНАЯ -----------------
@app.route('/')
def index():
    return render_template('index.html')

# ----------------- АВТОРИЗАЦИЯ -----------------
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'GET':
        return render_template('login.html')

    username = request.form.get('username', '').strip()
    password = request.form.get('password', '')

    user = User.query.filter_by(username=username).first()
    if not user or not user.check_password(password):
        flash('Неверный логин или пароль', 'danger')
        return redirect(url_for('login'))

    session['user_id'] = user.id
    flash('Добро пожаловать 👋', 'success')
    return redirect(url_for('programming_languages'))


@app.route('/logout')
def logout():
    session.pop('user_id', None)
    flash('Вы вышли из аккаунта', 'info')
    return redirect(url_for('index'))

# ----------------- РЕГИСТРАЦИЯ -----------------
@app.route('/register', methods=['GET', 'POST'])
def register():
    errors, form = {}, {}
    if request.method == 'GET':
        return render_template('register.html', errors=errors, form=form)

    form['username']  = request.form.get('username', '').strip()
    form['email']     = request.form.get('email', '').strip()
    form['full_name'] = request.form.get('full_name', '').strip()
    password          = request.form.get('password', '')
    confirm           = request.form.get('confirm', '')

    if not form['username']:
        errors['username'] = 'Введите логин'
    elif not USERNAME_RE.match(form['username']):
        errors['username'] = 'Логин 3–30 символов, латиница/цифры/_'
    elif User.query.filter_by(username=form['username']).first():
        errors['username'] = 'Такой логин уже занят'

    if not form['email']:
        errors['email'] = 'Введите e-mail'
    elif not EMAIL_RE.match(form['email']):
        errors['email'] = 'Некорректный e-mail'
    elif User.query.filter_by(email=form['email']).first():
        errors['email'] = 'Такой e-mail уже используется'

    if not PWD_RE.match(password):
        errors['password'] = 'Пароль слабый (8+ символов, буквы верх/низ и цифра)'
    if password != confirm:
        errors['confirm'] = 'Пароли не совпадают'

    if errors:
        return render_template('register.html', errors=errors, form=form)

    user = User(username=form['username'], email=form['email'], full_name=form['full_name'] or None)
    user.set_password(password)
    db.session.add(user)
    db.session.commit()

    flash('Регистрация успешна. Теперь войдите.', 'success')
    return redirect(url_for('login'))

# ----------------- ПРОФИЛЬ -----------------
@app.route('/profile')
def profile():
    uid = session.get('user_id')
    if not uid:
        return redirect(url_for('login'))

    user = User.query.get(uid)
    if not user:
        session.pop('user_id', None)
        return redirect(url_for('login'))

    materials_count = Material.query.count()
    theory_count    = Material.query.filter_by(type='theory').count()
    practice_count  = Material.query.filter_by(type='practice').count()

    by_lang = (db.session.query(Material.language, func.count(Material.id))
               .group_by(Material.language)
               .order_by(func.count(Material.id).desc())
               .all())

    recent = (MaterialOpen.query
              .filter_by(user_id=user.id)
              .order_by(MaterialOpen.opened_at.desc())
              .limit(10)
              .all())

    return render_template(
        'profile.html',
        user=user,
        materials_count=materials_count,
        theory_count=theory_count,
        practice_count=practice_count,
        by_lang=by_lang,
        recent=recent
    )

# ----------------- АДМИНКА -----------------
@app.route('/admin', methods=['GET', 'POST'])
def admin_dashboard():
    if request.method == 'POST':
        try:
            title    = request.form['title'].strip()
            mat_type = request.form['mat_type'].strip()
            language = request.form['language'].strip()
            file     = request.files['file']

            if not title or not file or file.filename == '':
                flash('Укажите название и выберите файл', 'warning')
                return redirect(url_for('admin_dashboard'))

            file_name = secure_filename(file.filename)
            file_data = file.read()

            m = Material(title=title, type=mat_type, language=language,
                         file_name=file_name, file_data=file_data)
            db.session.add(m)
            db.session.commit()
            flash('Материал добавлен', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Ошибка при добавлении: {e}', 'danger')

        return redirect(url_for('admin_dashboard'))

    materials = Material.query.order_by(Material.id.desc()).all()
    return render_template('admin_panel.html', materials=materials)


@app.route('/material/<int:material_id>/delete', methods=['POST'])
def delete_material(material_id):
    m = Material.query.get_or_404(material_id)
    try:
        db.session.delete(m)
        db.session.commit()
        flash('Материал удалён', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Ошибка удаления: {e}', 'danger')
    return redirect(url_for('admin_dashboard'))

# ----------------- МАТЕРИАЛЫ (СОРТИРОВКА + ПАГИНАЦИЯ) -----------------
def get_materials_query(sort, filter_type, search, language=None):
    query = Material.query
    if language:
        query = query.filter_by(language=language)
    if filter_type in ("theory", "practice"):
        query = query.filter_by(type=filter_type)
    if search:
        term = f"%{search}%"
        query = query.filter(or_(
            Material.title.ilike(term),
            Material.language.ilike(term),
            Material.type.ilike(term)
        ))
    if sort == "title":
        query = query.order_by(Material.title.asc())
    else:
        query = query.order_by(Material.created_at.desc(), Material.id.desc())
    return query

@app.route('/materials')
def materials():
    sort = request.args.get("sort", "date")
    filter_type = request.args.get("filter")  # theory / practice
    filter_lang = request.args.get("lang")    # python, js, cpp ...
    search = request.args.get("search", "").strip()
    page = request.args.get("page", 1, type=int)
    per_page = 10

    query = Material.query

    # фильтр по типу
    if filter_type in ("theory", "practice"):
        query = query.filter_by(type=filter_type)

    # фильтр по языку
    if filter_lang:
        query = query.filter(Material.language.ilike(filter_lang))

    # поиск (без учёта регистра)
    if search:
        term = f"%{search}%"
        query = query.filter(or_(
            Material.title.ilike(term),
            Material.language.ilike(term),
            Material.type.ilike(term)
        ))

    # сортировка
    if sort == "title":
        query = query.order_by(Material.title.asc())
    else:
        query = query.order_by(Material.created_at.desc(), Material.id.desc())

    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    items = pagination.items

    return render_template(
        "materials.html",
        materials=items,
        pagination=pagination,
        language=None,
        sort=sort,
        filter_type=filter_type,
        filter_lang=filter_lang,
        search=search
    )


@app.route('/materials/<language>')
def materials_by_language(language):
    sort = request.args.get("sort", "date")
    filter_type = request.args.get("filter")
    filter_lang = request.args.get("lang")
    search = request.args.get("search", "").strip()
    page = request.args.get("page", 1, type=int)
    per_page = 10

    query = Material.query.filter_by(language=language)

    if filter_type in ("theory", "practice"):
        query = query.filter_by(type=filter_type)

    if filter_lang:
        query = query.filter(Material.language.ilike(filter_lang))

    if search:
        term = f"%{search}%"
        query = query.filter(or_(
            Material.title.ilike(term),
            Material.language.ilike(term),
            Material.type.ilike(term)
        ))

    if sort == "title":
        query = query.order_by(Material.title.asc())
    else:
        query = query.order_by(Material.created_at.desc(), Material.id.desc())

    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    items = pagination.items

    return render_template(
        "materials.html",
        materials=items,
        pagination=pagination,
        language=language,
        sort=sort,
        filter_type=filter_type,
        filter_lang=filter_lang,
        search=search
    )


@app.route('/theory')
def theory_list():
    sort = request.args.get("sort", "date")
    filter_type = request.args.get("filter")
    search = request.args.get("search", "").strip()
    page = request.args.get("page", 1, type=int)
    per_page = 10

    query = Material.query.filter_by(type='theory')

    if search:
        term = f"%{search}%"
        query = query.filter(or_(
            Material.title.ilike(term),
            Material.language.ilike(term),
            Material.type.ilike(term)
        ))

    if sort == "title":
        query = query.order_by(Material.title.asc())
    else:
        query = query.order_by(Material.created_at.desc(), Material.id.desc())

    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    items = pagination.items

    return render_template(
        "materials.html",
        materials=items,
        pagination=pagination,
        language="theory",
        sort=sort,
        filter_type=filter_type,
        search=search
    )


@app.route('/practice')
def practice_list():
    sort = request.args.get("sort", "date")
    filter_type = request.args.get("filter")
    search = request.args.get("search", "").strip()
    page = request.args.get("page", 1, type=int)
    per_page = 10

    query = Material.query.filter_by(type='practice')

    if search:
        term = f"%{search}%"
        query = query.filter(or_(
            Material.title.ilike(term),
            Material.language.ilike(term),
            Material.type.ilike(term)
        ))

    if sort == "title":
        query = query.order_by(Material.title.asc())
    else:
        query = query.order_by(Material.created_at.desc(), Material.id.desc())

    pagination = query.paginate(page=page, per_page=per_page, error_out=False)
    items = pagination.items

    return render_template(
        "materials.html",
        materials=items,
        pagination=pagination,
        language="practice",
        sort=sort,
        filter_type=filter_type,
        search=search
    )


# ----------------- МАТЕРИАЛ -----------------
@app.route('/material/<int:material_id>')
def material_detail(material_id):
    m = Material.query.get_or_404(material_id)
    log_open(material_id)
    return render_template('material_detail.html', material=m)

@app.route('/material/<int:material_id>/view')
def material_view(material_id):
    m = Material.query.get_or_404(material_id)
    log_open(material_id)

    fname = (m.file_name or '').lower()
    mime, _ = guess_type(fname)

    if mime and (mime.startswith('image/') or mime == 'application/pdf' or mime.startswith('text/')):
        return send_file(io.BytesIO(m.file_data), mimetype=mime)

    if fname.endswith('.docx'):
        doc = Document(io.BytesIO(m.file_data))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        tables = [[ [cell.text.strip() for cell in row.cells] for row in t.rows] for t in doc.tables]
        return render_template('docx_view.html',
                               material=m,
                               title=m.title,
                               paragraphs=paragraphs,
                               tables=tables)

    flash('Этот тип файла нельзя показать онлайн. Скачайте его.', 'warning')
    return redirect(url_for('material_detail', material_id=material_id))

@app.route('/material/<int:material_id>/download')
def download_material(material_id):
    m = Material.query.get_or_404(material_id)
    log_open(material_id)
    return send_file(io.BytesIO(m.file_data),
                     as_attachment=True,
                     download_name=m.file_name)

# ----------------- ЯЗЫКИ -----------------
@app.route('/programming_languages')
def programming_languages():
    by_lang = (db.session.query(Material.language, func.count(Material.id))
               .group_by(Material.language)
               .order_by(func.count(Material.id).desc())
               .all())
    return render_template('programming_languages.html', by_lang=by_lang)

# ----------------- ОБРАТНАЯ СВЯЗЬ -----------------
@app.route("/contact", methods=["GET", "POST"])
def contact():
    if request.method == "POST":
        name = request.form.get("name")
        email = request.form.get("email")
        message_text = request.form.get("message")

        msg = Message(
            subject=f"Сообщение с сайта от {name}",
            recipients=[app.config["MAIL_USERNAME"]],
            body=f"От: {name} <{email}>\n\n{message_text}",
            sender=app.config["MAIL_DEFAULT_SENDER"]
        )

        try:
            mail.send(msg)
            flash("✅ Сообщение успешно отправлено!", "success")
        except Exception as e:
            flash(f"❌ Ошибка: {e}", "danger")

        return redirect(url_for("contact"))

    return render_template("contact.html")

# ----------------- ПОИСК -----------------
def highlight(text, q):
    """Подсветка совпадений (регистронезависимо)"""
    if not text or not q:
        return text
    pattern = re.compile(re.escape(q), re.IGNORECASE)
    return pattern.sub(lambda m: f"<mark>{m.group(0)}</mark>", text)

@app.route('/search')
def search():
    q = request.args.get('q', '').strip()
    page = request.args.get("page", 1, type=int)
    per_page = 10

    if not q:
        return render_template(
            'poisc.html',
            q=q,
            materials=[],
            total=0,
            pagination=None,
            message='Введите запрос в поле поиска'
        )

    term = f"%{q}%"
    query = (Material.query
             .filter(or_(
                 Material.title.ilike(term),
                 Material.file_name.ilike(term),
                 Material.language.ilike(term),
                 Material.type.ilike(term)
             )))

    pagination = query.order_by(Material.created_at.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )
    results = pagination.items

    # Подсветим найденное слово
    for r in results:
        r.title = highlight(r.title, q)
        r.language = highlight(r.language, q)
        r.type = highlight(r.type, q)

    return render_template(
        'poisc.html',
        q=q,
        materials=results,
        total=pagination.total,
        pagination=pagination,
        message=None
    )

# ----------------- ЗАПУСК -----------------
if __name__ == '__main__':
    app.run(debug=True)
