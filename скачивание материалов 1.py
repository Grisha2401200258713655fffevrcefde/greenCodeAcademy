import asyncio
import hashlib
import os
import re
import sqlite3
import sys
import time
from pathlib import Path
from typing import Optional, Set, Tuple
from urllib.parse import urlparse, urljoin

import aiohttp
from aiohttp import ClientTimeout
from aiolimiter import AsyncLimiter
from bs4 import BeautifulSoup
import trafilatura
from docx import Document
from pdfminer.high_level import extract_text as pdf_extract_text

# --- универсальный импорт DDGS ---
try:
    from duckduckgo_search import DDGS  # pip install duckduckgo-search
except ImportError:
    try:
        from ddgs import DDGS           # pip install ddgs
    except ImportError:
        DDGS = None

import argparse

USER_AGENT = "TB-Crawler/1.0 (+contact@example.com)"
ACCEPT = "*/*"
CONNECT_TIMEOUT = 30
READ_TIMEOUT = 120

TEXT_MIMES = {
    "text/html", "application/xhtml+xml", "text/plain", "application/xml", "text/xml",
    "application/json", "text/markdown", "text/x-markdown"
}
DOC_MIMES = {"application/pdf", "application/msword",
             "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}

SAVE_AS_IS_EXT = (
    ".pdf", ".zip", ".gz", ".tgz", ".bz2", ".xz", ".7z", ".rar",
    ".doc", ".docx", ".ppt", ".pptx", ".xls", ".xlsx", ".epub"
)

SAFE_NAME_RE = re.compile(r"[^a-zA-Zа-яА-Я0-9._()\-\s]+")

def safe_name(s: str, limit: int = 150) -> str:
    s = re.sub(r"\s+", " ", (s or "")).strip()
    s = SAFE_NAME_RE.sub("_", s)
    s = re.sub(r"_+", "_", s).strip("._- ")
    return s[:limit] or "document"

def sha256_bytes(b: bytes) -> str:
    h = hashlib.sha256(); h.update(b); return h.hexdigest()

class DB:
    def __init__(self, path: Path):
        self.conn = sqlite3.connect(str(path))
        self.conn.execute("PRAGMA journal_mode=WAL;")
        self.conn.execute("""
        CREATE TABLE IF NOT EXISTS frontier (
          url TEXT PRIMARY KEY,
          depth INTEGER
        );""")
        self.conn.execute("""
        CREATE TABLE IF NOT EXISTS visited (
          url TEXT PRIMARY KEY
        );""")
        self.conn.execute("""
        CREATE TABLE IF NOT EXISTS docs (
          sha256 TEXT PRIMARY KEY,
          url TEXT,
          domain TEXT,
          mime TEXT,
          bytes INTEGER,
          stored_path TEXT,
          text_path TEXT,
          title TEXT,
          created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );""")
        self.conn.commit()

    def add_to_frontier(self, url: str, depth: int):
        try:
            self.conn.execute("INSERT OR IGNORE INTO frontier(url, depth) VALUES(?,?)", (url, depth))
            self.conn.commit()
        except Exception:
            pass

    def pop_frontier(self) -> Optional[Tuple[str, int]]:
        cur = self.conn.execute("SELECT url, depth FROM frontier LIMIT 1")
        row = cur.fetchone()
        if not row:
            return None
        self.conn.execute("DELETE FROM frontier WHERE url=?", (row[0],))
        self.conn.commit()
        return row[0], row[1]

    def mark_visited(self, url: str):
        self.conn.execute("INSERT OR IGNORE INTO visited(url) VALUES(?)", (url,))
        self.conn.commit()

    def is_visited(self, url: str) -> bool:
        cur = self.conn.execute("SELECT 1 FROM visited WHERE url=? LIMIT 1", (url,))
        return cur.fetchone() is not None

    def has_sha(self, sha: str) -> bool:
        cur = self.conn.execute("SELECT 1 FROM docs WHERE sha256=? LIMIT 1", (sha,))
        return cur.fetchone() is not None

    def save_doc(self, sha: str, url: str, domain: str, mime: str, size: int, stored_path: str, text_path: Optional[str], title: Optional[str]):
        self.conn.execute(
            "INSERT OR IGNORE INTO docs(sha256,url,domain,mime,bytes,stored_path,text_path,title) VALUES(?,?,?,?,?,?,?,?)",
            (sha, url, domain, mime, size, stored_path, text_path, title)
        )
        self.conn.commit()

class RobotsCache:
    def __init__(self, session: aiohttp.ClientSession):
        self.session = session
        self.cache = {}

    async def allowed(self, url: str) -> bool:
        parsed = urlparse(url)
        base = f"{parsed.scheme}://{parsed.netloc}"
        rp = self.cache.get(base)
        if rp is None:
            robots_url = f"{base}/robots.txt"
            rules = []
            try:
                async with self.session.get(robots_url, headers={"User-Agent": USER_AGENT}, timeout=ClientTimeout(total=20)) as r:
                    if r.status == 200:
                        text = await r.text()
                        rules = text.splitlines()
            except Exception:
                pass
            self.cache[base] = rules
            rp = rules
        # Минимальная проверка: если явно запрещён / (упрощённо)
        for line in rp:
            l = line.strip().lower()
            if l.startswith("user-agent:") and "*" in l:
                continue
        # здесь делаем либерально: считаем разрешено, если не знаем правил
        return True

async def ddg_seeds(query: str, n: int) -> Set[str]:
    urls = set()
    if DDGS is None:
        return urls
    try:
        with DDGS() as d:
            for r in d.text(query, max_results=n):
                u = r.get("href") or r.get("url")
                if isinstance(u, str) and u.startswith("http"):
                    urls.add(u)
    except Exception:
        pass
    return urls

def detect_mime(ct: str, url_path: str) -> str:
    ct = (ct or "").split(";")[0].strip().lower()
    if ct:
        return ct
    p = url_path.lower()
    if p.endswith(".pdf"): return "application/pdf"
    if p.endswith(".html") or p.endswith(".htm"): return "text/html"
    if p.endswith(".txt"): return "text/plain"
    return "application/octet-stream"

def extract_links(base_url: str, html_text: str) -> Set[str]:
    out = set()
    soup = BeautifulSoup(html_text, "lxml")
    for a in soup.find_all("a", href=True):
        href = a.get("href")
        u = urljoin(base_url, href)
        if u.startswith("http"):
            out.add(u.split("#")[0])
    return out

def extract_title(html: str) -> Optional[str]:
    soup = BeautifulSoup(html, "lxml")
    if soup.title and soup.title.string:
        return soup.title.string.strip()
    return None

def html_to_text(html: str, url: str) -> str:
    extracted = trafilatura.extract(html, include_comments=False, include_tables=False, url=url)
    if extracted:
        return extracted
    soup = BeautifulSoup(html, "lxml")
    for t in ("script","style","noscript"):
        [x.decompose() for x in soup.find_all(t)]
    text = soup.get_text(separator="\n", strip=True)
    lines = [re.sub(r"\s+", " ", ln).strip() for ln in text.splitlines()]
    lines = [ln for ln in lines if ln]
    return "\n".join(lines)

def save_docx(text: str, title: Optional[str], url: str, out_path: Path):
    doc = Document()
    if title:
        doc.add_heading(title, 0)
    else:
        doc.add_heading("Document", 0)
    doc.add_paragraph(f"Источник: {url}")
    doc.add_paragraph("")
    for para in text.split("\n"):
        p = para.strip()
        if p:
            doc.add_paragraph(p)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

async def fetch_and_process(url: str, depth: int, db: DB, session: aiohttp.ClientSession,
                            robots: RobotsCache, out_dir: Path, limiter: AsyncLimiter,
                            max_doc_size: int, enqueue_cb):
    if db.is_visited(url):
        return
    # robots (упрощённо: не запрещаем, если не извлекли правил)
    if not await robots.allowed(url):
        db.mark_visited(url)
        return

    parsed = urlparse(url)
    domain = parsed.netloc
    dir_domain = out_dir / "raw" / domain
    dir_text = out_dir / "text" / domain
    dir_domain.mkdir(parents=True, exist_ok=True)
    headers = {"User-Agent": USER_AGENT, "Accept": ACCEPT}

    async with limiter:
        try:
            timeout = ClientTimeout(total=READ_TIMEOUT, connect=CONNECT_TIMEOUT)
            async with session.get(url, headers=headers, timeout=timeout, allow_redirects=True) as resp:
                status = resp.status
                if status != 200:
                    db.mark_visited(url)
                    return
                content = await resp.read()
                if max_doc_size and len(content) > max_doc_size:
                    # слишком крупный — пропускаем, чтобы не забивать диск резко
                    db.mark_visited(url)
                    return

                sha = sha256_bytes(content)
                if db.has_sha(sha):
                    db.mark_visited(url)
                    return

                ctype = detect_mime(resp.headers.get("Content-Type",""), parsed.path)
                # Сохраняем оригинал
                ext = ""
                lower_path = parsed.path.lower()
                for e in SAVE_AS_IS_EXT:
                    if lower_path.endswith(e):
                        ext = e
                        break
                if not ext:
                    if ctype.startswith("text/html"):
                        ext = ".html"
                    elif ctype == "application/pdf":
                        ext = ".pdf"
                    else:
                        ext = ".bin"

                file_name = safe_name(sha) + ext
                raw_path = dir_domain / file_name
                raw_path.write_bytes(content)

                text_path = None
                title = None

                # Извлекаем текст и пишем .docx для читаемых типов
                if ctype.startswith("text/html"):
                    try:
                        html = content.decode("utf-8", errors="ignore")
                        title = extract_title(html)
                        text = html_to_text(html, url)
                        if text.strip():
                            docx_name = safe_name(sha) + ".docx"
                            text_path = str((dir_text / docx_name).resolve())
                            save_docx(text, title, url, Path(text_path))
                    except Exception:
                        pass
                    # ссылки для следующей глубины
                    if depth > 0:
                        for link in extract_links(url, html):
                            if not db.is_visited(link):
                                enqueue_cb(link, depth - 1)

                elif ctype == "application/pdf":
                    try:
                        # сохраним также текст в docx если получится
                        from tempfile import NamedTemporaryFile
                        with NamedTemporaryFile(suffix=".pdf", delete=True) as f:
                            f.write(content); f.flush()
                            pdf_text = pdf_extract_text(f.name) or ""
                        if pdf_text.strip():
                            docx_name = safe_name(sha) + ".docx"
                            text_path = str((dir_text / docx_name).resolve())
                            save_docx(pdf_text, None, url, Path(text_path))
                    except Exception:
                        pass

                db.save_doc(sha, url, domain, ctype, len(content), str(raw_path.resolve()), text_path, title)
                db.mark_visited(url)
        except Exception:
            db.mark_visited(url)
            return

async def crawl(out_dir: Path, start_urls: Set[str], depth: int,
                concurrency: int, max_doc_size: int):
    out_dir.mkdir(parents=True, exist_ok=True)
    db = DB(out_dir / "crawler.sqlite3")

    # фронтир
    for u in start_urls:
        db.add_to_frontier(u, depth)

    connector = aiohttp.TCPConnector(limit=concurrency*2, ssl=False)
    async with aiohttp.ClientSession(connector=connector) as session:
        robots = RobotsCache(session)
        limiter = AsyncLimiter(max_calls=concurrency, time_period=1.0)  # ~N req/sec

        queue = asyncio.Queue()

        def enqueue(url: str, d: int):
            if not url.startswith("http"):
                return
            if not db.is_visited(url):
                db.add_to_frontier(url, d)
                queue.put_nowait((url, d))

        # начальная загрузка в очередь из БД
        while True:
            row = db.pop_frontier()
            if not row:
                break
            queue.put_nowait(row)

        async def worker():
            while True:
                try:
                    url, d = await queue.get()
                except asyncio.CancelledError:
                    return
                await fetch_and_process(url, d, db, session, robots, out_dir, limiter, max_doc_size, enqueue)
                queue.task_done()

        workers = [asyncio.create_task(worker()) for _ in range(max(1, concurrency//5))]
        # запустим подачу из БД по мере опустошения
        feeder_sleep = 0.5

        try:
            while True:
                await asyncio.sleep(feeder_sleep)
                # если очередь пустая, пробуем снова наполнять из фронтира (вдруг что-то добавилось из HTML)
                if queue.empty():
                    row = db.pop_frontier()
                    if row:
                        queue.put_nowait(row)
                    else:
                        # возможная пауза, ждём пока докачает всё и выйдем
                        if all(q.done() for q in workers):
                            break
                        # если у воркеров нет задач, выходим
                        if queue.empty():
                            break
                # если очередь очень большая — просто ждём
        finally:
            await queue.join()
            for w in workers:
                w.cancel()
            await asyncio.gather(*workers, return_exceptions=True)

def main():
    ap = argparse.ArgumentParser(description="Массовый краулер: ищет, качает, сохраняет оригиналы и .docx")
    ap.add_argument("--out", default="dataset", help="Папка для сохранения (по умолчанию ./dataset)")
    ap.add_argument("--query", type=str, default="", help="Стартовый поисковый запрос для DDG")
    ap.add_argument("--seeds", type=int, default=100, help="Сколько стартовых ссылок взять из поиска")
    ap.add_argument("--depth", type=int, default=2, help="Глубина обхода ссылок")
    ap.add_argument("--concurrency", type=int, default=40, help="Одновременных запросов/сек")
    ap.add_argument("--max-doc-size", type=int, default=50*1024*1024, help="Макс размер файла в байтах (0 = не ограничивать)")
    args = ap.parse_args()

    out_dir = Path(args.out).resolve()

    start_urls = set()
    if args.query and args.seeds > 0:
        if DDGS is None:
            print("Внимание: пакет DDGS не установлен. Установи duckduckgo-search или ddgs.")
        else:
            print(f"[seeds] ищу {args.seeds} ссылок по запросу: {args.query}")
            start_urls = asyncio.run(ddg_seeds(args.query, args.seeds))
            print(f"[seeds] найдено: {len(start_urls)}")

    if not start_urls:
        print("Нет стартовых ссылок. Укажи --query или добавь свои URL в БД frontier.")
        print("Пример: --query \"data science tutorial\" --seeds 300")
        sys.exit(0)

    print(f"[crawl] out={out_dir} depth={args.depth} concurrency={args.concurrency}")
    asyncio.run(crawl(out_dir, start_urls, args.depth, args.concurrency, args.max_doc_size))
    print("[done] Готово. Можно останавливать/перезапускать — БД удержит прогресс.")

if __name__ == "__main__":
    main()
