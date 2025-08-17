# file: materials_downloader_docx.py

import os
import re
import requests
from bs4 import BeautifulSoup
import feedparser
from docx import Document

LANGUAGES = {
    "python": ["python"],
    "javascript": ["javascript", "js", "node"],
    "java": ["java"],
    "c++": ["c++", "cpp"],
    "go": ["go", "golang"],
    "rust": ["rust"],
}
TOPICS = {
    "syntax": ["syntax", "синтаксис", "операторы", "variables", "переменные", "statements"],
    "oop": ["oop", "object", "объект", "class", "класс", "инкапсуляция"],
    "libraries": ["library", "библиотека", "package", "модуль", "framework", "фреймворк"],
    "web": ["web", "http", "api", "интернет", "django", "flask", "express"],
    "data": ["data", "данные", "csv", "json", "парсинг", "парсер"],
    "testing": ["test", "тест", "pytest", "unit"],
    "tools": ["tool", "инструмент", "cli", "pip", "npm", "maven"],
    "misc": [],
}

def get_devto_rss(language: str):
    url = f"https://dev.to/feed/tag/{language}"
    feed = feedparser.parse(url)
    articles = []
    for entry in feed.entries[:5]:
        articles.append({
            "title": entry.title,
            "link": entry.link
        })
    return articles

def download_article(url: str) -> str:
    try:
        resp = requests.get(url, timeout=10)
        soup = BeautifulSoup(resp.content, "html.parser")
        article = soup.find("div", {"class": "crayons-article__main"})
        if not article:
            article = soup.find("article")
        if article:
            for tag in article(["script", "style"]):
                tag.decompose()
            text = article.get_text(separator="\n", strip=True)
            return text
        return ""
    except Exception as e:
        print(f"Ошибка при загрузке {url}: {e}")
        return ""

def classify_topic(content: str) -> str:
    content_lower = content.lower()
    for topic, keywords in TOPICS.items():
        if any(kw in content_lower for kw in keywords):
            return topic
    return "misc"

def slugify(text):
    text = text.lower()
    text = re.sub(r'[^a-z0-9а-яё]+', '-', text)
    return text.strip('-')

def save_article_docx(base_path: str, lang: str, topic: str, title: str, content: str):
    # Формируем путь: <base_path>/<lang>/<topic>/
    dir_path = os.path.join(base_path, lang, topic)
    os.makedirs(dir_path, exist_ok=True)
    filename = slugify(title)[:50] + ".docx"
    file_path = os.path.join(dir_path, filename)
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    doc.save(file_path)

def main():
    print("Введите ПОЛНЫЙ путь к папке, куда сохранять материалы (например, D:\\Materials или /home/user/materials):")
    base_path = input("> ").strip()
    if not base_path:
        print("Путь не задан. Завершение работы.")
        return

    for lang, tags in LANGUAGES.items():
        print(f"--- Язык: {lang} ---")
        for tag in tags:
            articles = get_devto_rss(tag)
            for art in articles:
                print(f"  → {art['title']}")
                content = download_article(art['link'])
                if not content or len(content) < 100:
                    continue
                topic = classify_topic(content)
                save_article_docx(base_path, lang, topic, art['title'], content)
    print(f"Готово! Все материалы скачаны и разложены по папкам в: {base_path}")

if __name__ == "__main__":
    main()
